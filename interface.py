import os
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from PyPDF2 import PdfReader
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# File Paths
DATA_FILE_PATH = os.path.join(os.path.dirname(__file__), "program_data")
EXAMPLE_DOCS_FILE = os.path.join(DATA_FILE_PATH, "example_docs.json")
INFORMATION_DOCS_FILE = os.path.join(DATA_FILE_PATH, "information_docs.json")
SETTINGS_FILE = os.path.join(DATA_FILE_PATH, "settings.json")
DEFAULT_DOWNLOAD_DIR = os.path.join(DATA_FILE_PATH, "generated_docs")

# Global variable dictionaries to store uploaded document text
example_doc_dict = {}
information_doc_dict = {}

def load_settings():
    """Load settings from the settings file or set defaults."""
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r") as f:
            settings = json.load(f)
    else:
        settings = {"download_path": DEFAULT_DOWNLOAD_DIR, "file_type": ".docx"}
        save_settings(settings)
    return settings

def save_settings(settings):
    """Save settings to the settings file."""
    with open(SETTINGS_FILE, "w") as f:
        json.dump(settings, f)

def save_docs():
    """Save both example and information documents to file."""
    save_example_docs()
    save_information_docs()

def load_docs():
    """Load both example and information documents from file."""
    load_example_docs()
    load_information_docs()

def save_example_docs():
    """Save example documents to file."""
    global example_doc_dict
    with open(EXAMPLE_DOCS_FILE, "w") as f:
        json.dump(example_doc_dict, f)

def load_example_docs():
    """Load example documents from file."""
    if os.path.exists(EXAMPLE_DOCS_FILE):
        global example_doc_dict
        with open(EXAMPLE_DOCS_FILE, "r") as f:
            example_doc_dict = json.load(f)

def save_information_docs():
    """Save information documents to file."""
    global information_doc_dict
    with open(INFORMATION_DOCS_FILE, "w") as f:
        json.dump(information_doc_dict, f)

def load_information_docs():
    """Load information documents from file."""
    if os.path.exists(INFORMATION_DOCS_FILE):
        global information_doc_dict
        with open(INFORMATION_DOCS_FILE, "r") as f:
            information_doc_dict = json.load(f)

def show_uploaded_files():
    """Show a popup with the filenames of all uploaded documents."""
    global example_doc_dict, information_doc_dict
    popup = tk.Toplevel(root)
    popup.title("Uploaded Files")

    def delete_selected_from_listbox(listbox, doc_dict):
        """Delete the selected file(s) from the provided listbox and dictionary."""
        selected_items = listbox.curselection()
        for i in reversed(selected_items): 
            filename = listbox.get(i)
            del doc_dict[filename]
            listbox.delete(i)
        save_docs()

    # Example Documents
    tk.Label(popup, text="Example Documents:").pack(anchor="center", padx=10, pady=5)
    example_listbox = tk.Listbox(popup, width=50, selectmode=tk.MULTIPLE)
    example_listbox.pack(padx=10, pady=5)
    for filename in example_doc_dict.keys():
        example_listbox.insert(tk.END, filename)
    example_button_frame = tk.Frame(popup)
    example_button_frame.pack(pady=5)
    tk.Button(
        example_button_frame, text="Delete Selected", 
        command=lambda: delete_selected_from_listbox(example_listbox, example_doc_dict)
    ).pack(side=tk.LEFT, padx=5)
    tk.Button(example_button_frame, text="Clear All", command=clear_examples).pack(side=tk.LEFT, padx=5)

    # Informational Documents
    tk.Label(popup, text="Informational Documents:").pack(anchor="center", padx=10, pady=5)
    info_listbox = tk.Listbox(popup, width=50, selectmode=tk.MULTIPLE)
    info_listbox.pack(padx=10, pady=5)
    for filename in information_doc_dict.keys():
        info_listbox.insert(tk.END, filename)
    info_button_frame = tk.Frame(popup)
    info_button_frame.pack(pady=5)
    tk.Button(
        info_button_frame, text="Delete Selected", 
        command=lambda: delete_selected_from_listbox(info_listbox, information_doc_dict)
    ).pack(side=tk.LEFT, padx=5)
    tk.Button(info_button_frame, text="Clear All", command=clear_information).pack(side=tk.LEFT, padx=5)

    # Close Button
    tk.Button(popup, text="Close", command=popup.destroy).pack(pady=10)

def clear_examples():
    """Clear uploaded document data."""
    global example_doc_dict
    response = messagebox.askyesno("Clear Documents", "Are you sure you want to remove all documents?")
    if response:
        example_doc_dict = {}
        if os.path.exists(EXAMPLE_DOCS_FILE):
            save_example_docs()
        messagebox.showinfo("Success", "All documents have been removed.")

def clear_information():
    """Clear uploaded information documents."""
    global information_doc_dict
    response = messagebox.askyesno("Clear Documents", "Are you sure you want to remove all documents?")
    if response:
        information_doc_dict = {}
        if os.path.exists(INFORMATION_DOCS_FILE):
            save_information_docs()
        messagebox.showinfo("Success", "All documents have been removed.")

def extract_text_from_document(filepath):
    """Extract text from the document."""
    if filepath.endswith(".docx"):
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    elif filepath.endswith(".pdf"):
        try:
            reader = PdfReader(filepath)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            messagebox.showerror("Error", f"Failed to extract text from PDF: {e}")
            return ""
    elif filepath.endswith(".txt"):
        with open(filepath, "r") as f:
            return f.read()
    else:
        return ""
    
def upload_documents(doc_dict, type):
    """Upload and store documents in the given dictionary."""
    filepaths = filedialog.askopenfilenames(
        title=f"Upload {type} Documents",
        filetypes=[("Word, PDF, and Text Documents", "*.docx *.pdf *.txt")],
    )

    if not filepaths:
        return

    for filepath in filepaths:
        filename = os.path.basename(filepath)
        if filename in doc_dict:
            response = messagebox.askyesno("File Exists", f"{filename} already exists. Replace it?")
            if not response:
                continue

        # Extract and store document text
        doc_dict[filename] = extract_text_from_document(filepath)
        print(f"Document uploaded: {filename}")

    messagebox.showinfo("Success", f"All {type} documents have been uploaded.")

def upload_example_docs():
    """Upload example documents."""
    global example_doc_dict
    upload_documents(example_doc_dict, "Example")
    save_example_docs()

def upload_information_docs():
    """Upload informational documents."""
    global information_doc_dict
    upload_documents(information_doc_dict, "Informational")
    save_information_docs()

def set_download_path(event):
    folderpath = filedialog.askdirectory(title="Select Download Path")
    if folderpath:
        download_path_entry.config(state="normal")
        download_path_entry.delete(0, tk.END)
        download_path_entry.insert(0, folderpath)
        download_path_entry.config(state="readonly")
        settings["download_path"] = folderpath
        save_settings(settings)

def get_data_from_ai(formatting, information):
    """Placeholder for AI logic.""" 
    final_document_text = f"Formatting:\n{formatting}\n\n"
    final_document_text += f"Information:\n{information}\n\n"

    return "Sample AI Document\n\n" + final_document_text

def generate_document():
    """Generate the final document."""
    global example_doc_dict, information_doc_dict
    formatting = formatting_textbox.get("1.0", tk.END).strip()
    information = information_textbox.get("1.0", tk.END).strip()
    download_path = download_path_entry.get().strip()
    filename = filename_entry.get()

    # Append text from uploaded documents
    for name, text in example_doc_dict.items():
        formatting += f"\n\nText from {name}:\n\n{text}"

    for name, text in information_doc_dict.items():
        information += f"\n\nText from {name}:\n\n{text}"

    if not formatting or not information or not download_path or not filename:
        messagebox.showerror("Error", "Please fill in all fields and set a download path and file name.")
        return

    text_string = get_data_from_ai(formatting, information)
    output_filepath = f"{download_path}\{filename}{file_type.get()}" # test different file types
    output_doc = None

    if file_type.get() == ".docx":
        output_doc = Document()
        output_doc.add_paragraph(text_string)
        output_doc.save(output_filepath)
    elif file_type.get() == ".pdf":
        try:
            # Create the PDF canvas
            c = canvas.Canvas(output_filepath, pagesize=letter)
            text_object = c.beginText(50, 750)  # Starting position on the page (x, y)

            # Wrap text to fit within the page
            lines = text_string.splitlines()
            for line in lines:
                text_object.textLine(line)
                if text_object.getY() < 50:
                    c.drawText(text_object)
                    c.showPage()
                    text_object = c.beginText(50, 750)

            # Save the last page
            c.drawText(text_object)
            c.save()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate PDF: {e}")
    elif file_type.get() == ".txt":
        output_doc = open(output_filepath, "w")
        output_doc.write(text_string)
        output_doc.close()

    messagebox.showinfo("Success", f"Document generated and saved to {output_filepath}")
    
if not os.path.exists(DEFAULT_DOWNLOAD_DIR):
    os.makedirs(DEFAULT_DOWNLOAD_DIR)
load_docs()
settings = load_settings()
default_download_path = settings.get("download_path", DEFAULT_DOWNLOAD_DIR)

# Create the main window
root = tk.Tk()
root.title("Document Generator")
root.geometry("600x700")

# Formatting Section
tk.Label(root, text="Describe The formatting of the document and upload examples of the format:").pack(anchor="center", padx=10, pady=5)
formatting_textbox = tk.Text(root, height=5, width=70)
formatting_textbox.pack(padx=10, pady=5)
tk.Button(root, text="Upload Example Documents", command=upload_example_docs).pack(pady=5)

# Information Section
tk.Label(root, text="Paste and upload any information needed for the document here:").pack(anchor="center", padx=10, pady=5)
information_textbox = tk.Text(root, height=5, width=70)
information_textbox.pack(padx=10, pady=5)
tk.Button(root, text="Upload Informational Documents", command=upload_information_docs).pack(pady=5)

# Download Path Section
tk.Label(root, text="Set File Download Location:").pack(anchor="center", padx=10, pady=5)
download_path_entry = tk.Entry(root, width=50)
download_path_entry.pack(pady=5)
download_path_entry.insert(0, default_download_path)
download_path_entry.bind("<Button-1>", set_download_path)
download_path_entry.config(state="readonly")

# Filename Section
tk.Label(root, text="Set New File Name:").pack(anchor="center", padx=10, pady=5)
filename_entry = tk.Entry(root, width=50)
filename_entry.pack(pady=5)

# File Type Section
tk.Label(root, text="Select File Type:").pack(anchor="center", padx=10, pady=5)
file_type_options = [".docx", ".pdf", ".txt"]
file_type = tk.StringVar(value=settings["file_type"])

def update_file_type(*args):
    settings["file_type"] = file_type.get()
    save_settings(settings)

file_type.trace("w", update_file_type)
file_type_dropdown = tk.OptionMenu(root, file_type, *file_type_options)
file_type_dropdown.pack(pady=5)

# Generate Button
tk.Button(root, text="Generate", command=generate_document).pack(pady=20)

# Show Uploaded Files Button
tk.Button(root, text="Show Uploaded Documents", command=show_uploaded_files).pack(pady=5)

root.mainloop()