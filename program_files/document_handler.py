import os
import json
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from PyPDF2 import PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from generator import get_data_from_ai
from config import EXAMPLE_DOCS_FILE, INFORMATION_DOCS_FILE

example_doc_dict = {}
information_doc_dict = {}

def get_example_docs():
    """Return the example documents dictionary."""
    global example_doc_dict
    return example_doc_dict

def get_information_docs():
    """Return the information documents dictionary."""
    global information_doc_dict
    return information_doc_dict

def save_docs():
    """Save both example and information documents to file."""
    save_example_docs()
    save_information_docs()

def load_docs():
    """Load both example and information documents from file."""
    load_example_docs()
    load_information_docs()

def save_example_docs():
    """Save example documents to file."""
    global example_doc_dict
    with open(EXAMPLE_DOCS_FILE, "w") as f:
        json.dump(example_doc_dict, f)

def load_example_docs():
    """Load example documents from file."""
    if os.path.exists(EXAMPLE_DOCS_FILE):
        global example_doc_dict
        with open(EXAMPLE_DOCS_FILE, "r") as f:
            example_doc_dict = json.load(f)

def save_information_docs():
    """Save information documents to file."""
    global information_doc_dict
    with open(INFORMATION_DOCS_FILE, "w") as f:
        json.dump(information_doc_dict, f)

def load_information_docs():
    """Load information documents from file."""
    if os.path.exists(INFORMATION_DOCS_FILE):
        global information_doc_dict
        with open(INFORMATION_DOCS_FILE, "r") as f:
            information_doc_dict = json.load(f)

def clear_example(filename):
    global example_doc_dict
    if filename in example_doc_dict:
        del example_doc_dict[filename]
        save_example_docs()
        messagebox.showinfo("Success", f"{filename} has been removed.")
        return example_doc_dict

def clear_all_examples():
    """Clear uploaded document data."""
    global example_doc_dict
    response = messagebox.askyesno("Clear Documents", "Are you sure you want to remove all documents?")
    if response:
        example_doc_dict = {}
        if os.path.exists(EXAMPLE_DOCS_FILE):
            save_example_docs()
        messagebox.showinfo("Success", "All documents have been removed.")
        return example_doc_dict

def clear_information(filename):
    global information_doc_dict
    if filename in information_doc_dict:
        del information_doc_dict[filename]
        save_information_docs()
        messagebox.showinfo("Success", f"{filename} has been removed.")
        return information_doc_dict

def clear_all_information():
    """Clear uploaded information documents."""
    global information_doc_dict
    response = messagebox.askyesno("Clear Documents", "Are you sure you want to remove all documents?")
    if response:
        information_doc_dict = {}
        if os.path.exists(INFORMATION_DOCS_FILE):
            save_information_docs()
        messagebox.showinfo("Success", "All documents have been removed.")
        return information_doc_dict

def extract_text_from_document(filepath):
    """Extract text from the document."""
    if filepath.endswith(".docx"):
        doc = Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    elif filepath.endswith(".pdf"):
        try:
            reader = PdfReader(filepath)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            messagebox.showerror("Error", f"Failed to extract text from PDF: {e}")
            return ""
    elif filepath.endswith(".txt"):
        with open(filepath, "r") as f:
            return f.read()
    else:
        return ""
    
def upload_documents(doc_dict, type):
    """Upload and store documents in the given dictionary."""
    filepaths = filedialog.askopenfilenames(
        title=f"Upload {type} Documents",
        filetypes=[("Word, PDF, and Text Documents", "*.docx *.pdf *.txt")],
    )

    if not filepaths:
        return

    for filepath in filepaths:
        filename = os.path.basename(filepath)
        if filename in doc_dict:
            response = messagebox.askyesno("File Exists", f"{filename} already exists. Replace it?")
            if not response:
                continue

        # Extract and store document text
        doc_dict[filename] = extract_text_from_document(filepath)
        print(f"Document uploaded: {filename}")

    messagebox.showinfo("Success", f"{type} documents have been uploaded.")

def upload_example_docs():
    """Upload example documents."""
    global example_doc_dict
    upload_documents(example_doc_dict, "Example")
    save_example_docs()

def upload_information_docs():
    """Upload informational documents."""
    global information_doc_dict
    upload_documents(information_doc_dict, "Informational")
    save_information_docs()

def generate_document(formatting_textbox, information_textbox, download_path_entry, filename_entry, file_type):
    """Generate the final document."""
    global example_doc_dict, information_doc_dict

    formatting_description = formatting_textbox.get('1.0', tk.END).strip()
    formatting_examples = ""
    information = information_textbox.get('1.0', tk.END).strip()
    download_path = download_path_entry.get().strip()
    filename = filename_entry.get()

    # add textbox text to the saved data
    if formatting_description:
        example_doc_dict['textbox'] = formatting_description
    else:
        example_doc_dict['textbox'] = ""

    if information:
        information_doc_dict['textbox'] = information
    else:
        information_doc_dict['textbox'] = ""

    save_docs()

    # Append text from uploaded documents
    for name, text in example_doc_dict.items():
        if name == "textbox":
            continue
        formatting_examples += f"\n\nText from {name}:\n\n{text}"

    for name, text in information_doc_dict.items():
        if name == "textbox":
            continue
        information += f"\n\nText from {name}:\n\n{text}"

    if not formatting_examples and formatting_description == "":
        messagebox.showerror("Error", "Please describe the formatting of the document.")
        return
    elif not information:
        messagebox.showerror("Error", "Please provide the information for the document.")
        return
    elif not download_path:
        messagebox.showerror("Error", "Please set a download path.")
        return
    elif not filename:
        messagebox.showerror("Error", "Please set a file name.")
        return

    try:
        # Generate the document text using AI
        text_string = get_data_from_ai(formatting_description, formatting_examples, information)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate document: {e}")
        return
    
    output_filepath = f"{download_path}\{filename}{file_type.get()}"
    output_doc = None

    if file_type.get() == ".docx":
        output_doc = Document()
        output_doc.add_paragraph(text_string)
        output_doc.save(output_filepath)
    elif file_type.get() == ".pdf":
        try:
            # Create the PDF canvas
            c = canvas.Canvas(output_filepath, pagesize=letter)
            text_object = c.beginText(50, 750)  # Starting position on the page (x, y)

            # Wrap text to fit within the page
            lines = text_string.splitlines()
            for line in lines:
                text_object.textLine(line)
                if text_object.getY() < 50:
                    c.drawText(text_object)
                    c.showPage()
                    text_object = c.beginText(50, 750)

            # Save the last page
            c.drawText(text_object)
            c.save()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate PDF: {e}")
    elif file_type.get() == ".txt":
        output_doc = open(output_filepath, "w")
        output_doc.write(text_string)
        output_doc.close()

    messagebox.showinfo("Success", f"Document generated and saved to {output_filepath}")